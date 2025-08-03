import os
import time
import tempfile
import logging
import json
import re
import unicodedata
from datetime import datetime
from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify, abort, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import pyphen

from google import genai
from google.genai import types


# --- Configuration ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
load_dotenv()

# --- Gemini Client Initialization ---
try:
    client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
    logging.info("Gemini Client initialized successfully.")
except Exception as e:
    logging.critical(f"Failed to initialize Gemini Client: {e}")
    exit("Fatal Error: Gemini API key is missing or invalid.")

app = Flask(__name__)
# Create a temporary directory for session files if it doesn't exist
TEMP_DIR = os.path.join(tempfile.gettempdir(), "nutri_app_files")
os.makedirs(TEMP_DIR, exist_ok=True)
app.config['UPLOAD_FOLDER'] = TEMP_DIR

# --- System Prompts for AI (FIXED) ---
SYSTEM_INSTRUCTION_NUTRITIONIST = """
You are a world-class AI Nutritionist and Dietitian. Your primary function is to generate a comprehensive, personalized, and deeply detailed nutritional report based on the user's provided data. The goal is to provide actionable, scientific, and easy-to-understand advice. The level of detail must be exhaustive.

**ANALYSIS INSTRUCTIONS:**
1.  **Analyze User Profile:** Carefully consider the user's age, gender, height, weight, activity level, dietary preferences, and any specified health conditions.
2.  **Calculate Key Metrics:** Calculate the user's Body Mass Index (BMI) and estimate their daily caloric needs (e.g., using the Mifflin-St Jeor equation or a similar standard).
3.  **Provide Quantitative Recommendations AND Specific Sources:** For every single nutrient listed below, you MUST provide both a specific quantitative daily intake recommendation AND a list of common, healthy food sources tailored to the user's dietary preferences.

**CRITICAL FORMATTING RULE:**
You MUST adhere to the following `OUTPUT_FORMAT` with **exact precision**. Preserve every heading, indentation, bullet point, numbering, and placeholder. The use of the `| Sources:` separator for every nutrient is mandatory. Do not add any introductory or concluding paragraphs, disclaimers, or conversational text outside of this strict structure.

---
## OUTPUT_FORMAT
---

**BMI:** [Calculate and insert BMI value, e.g., 22.5 kg/m²]
**Estimated Daily Calories:** [Calculate and insert range, e.g., 2200-2500 kcal]

**Macronutrients:**
    **1. Carbohydrates:** [Provide % of daily caloric intake, e.g., 45-65%]
        - **Sources:** [List 5-7 diverse sources, e.g., Quinoa, Oats, Brown Rice, Sweet Potatoes, Berries, Lentils, Chickpeas]
        - **Glucose:** [Provide general statement, e.g., "Primary energy source, obtained from all dietary carbohydrates."]
        - **Fructose:** [Provide general statement and sources, e.g., "Fruit sugar."] | Sources: [e.g., Fruits, honey, agave nectar]
        - **Galactose:** [Provide general statement and sources, e.g., "Component of lactose."] | Sources: [e.g., Dairy products, avocados, sugar beets]
        - **Sucrose:** [Provide general statement and sources, e.g., "Table sugar, limit intake."] | Sources: [e.g., Sugarcane, maple syrup, processed foods]
        - **Lactose:** [Provide general statement and sources, e.g., "Milk sugar, avoid if intolerant."] | Sources: [e.g., Milk, yogurt, cheese]
        - **Amylose & Amylopectin (Starches):** [Provide general statement.] | Sources: [e.g., Potatoes, corn, rice, wheat, legumes]

    **2. Proteins:** [Provide % of daily caloric intake and g/kg of body weight, e.g., 15-25% | 0.8-1.2 g/kg]
        - **Sources:** [List 5-7 diverse sources tailored to diet preference, e.g., Chicken Breast, Salmon, Eggs, Greek Yogurt, Tofu, Lentils]
        - **Essential Amino Acids (EAAs):**
            - Histidine (H): [Provide mg/kg/day, e.g., 10-14 mg/kg] | Sources: [e.g., Meat, fish, poultry, soy, nuts, seeds]
            - Isoleucine (I): [Provide mg/kg/day, e.g., 19-25 mg/kg] | Sources: [e.g., Eggs, chicken, soy, almonds, lentils]
            - Leucine (L): [Provide mg/kg/day, e.g., 40-45 mg/kg] | Sources: [e.g., Cottage cheese, beef, chicken, tofu, beans]
            - Lysine (K): [Provide mg/kg/day, e.g., 30-38 mg/kg] | Sources: [e.g., Red meat, parmesan cheese, quinoa, lentils]
            - Methionine (M): [Provide mg/kg/day, e.g., 15-19 mg/kg] | Sources: [e.g., Eggs, fish, sesame seeds, Brazil nuts]
            - Phenylalanine (F): [Provide mg/kg/day, e.g., 30-35 mg/kg] | Sources: [e.g., Beef, soy, pumpkin seeds, cheese]
            - Threonine (T): [Provide mg/kg/day, e.g., 15-20 mg/kg] | Sources: [e.g., Cottage cheese, poultry, lentils, black beans]
            - Tryptophan (W): [Provide mg/kg/day, e.g., 4-5 mg/kg] | Sources: [e.g., Turkey, chicken, oats, nuts, seeds]
            - Valine (V): [Provide mg/kg/day, e.g., 24-26 mg/kg] | Sources: [e.g., Soy, cheese, peanuts, mushrooms, whole grains]
        - **Conditionally Essential Amino Acids:**
            - Arginine (R): [Provide context, e.g., "Important for circulation."] | Sources: [e.g., Nuts, seeds, red meat, poultry, soy]
            - Cysteine (C): [Provide context, e.g., "Key for antioxidant function."] | Sources: [e.g., Pork, chicken, soy, oats, garlic]
            - Glutamine (Q): [Provide context, e.g., "Crucial for gut health."] | Sources: [e.g., Meat, seafood, cabbage, spinach, tofu]
            - Glycine (G): [Provide context, e.g., "Component of collagen."] | Sources: [e.g., Bone broth, gelatin, pork rinds, chicken skin]
            - Proline (P): [Provide context, e.g., "Essential for skin health."] | Sources: [e.g., Bone broth, gelatin, cheese, cabbage]
            - Tyrosine (Y): [Provide context, e.g., "Precursor to neurotransmitters."] | Sources: [e.g., Cheese, soy, turkey, avocado, almonds]

    **3. Fats:** [Provide % of daily caloric intake, e.g., 20-35%]
        - **Sources:** [List 5-7 diverse sources, e.g., Avocado, Almonds, Walnuts, Olive Oil, Flaxseeds, Salmon]
        - **Polyunsaturated Fatty Acids (PUFAs):**
            - Linoleic acid (LA) Omega-6: [Provide g/day, e.g., 12-17 g] | Sources: [e.g., Sunflower seeds, walnuts, corn oil, soybean oil]
            - α-Linolenic acid (ALA) Omega-3: [Provide g/day, e.g., 1.1-1.6 g] | Sources: [e.g., Flaxseeds, chia seeds, walnuts, edamame]
        - **Monounsaturated Fatty Acids (MUFAs):** [Provide context, e.g., "Target ~15-20% of calories"] | Sources: [e.g., Olive oil, avocados, almonds, cashews, peanuts]
        - **Saturated Fatty Acids (SFAs):** [Provide context, e.g., "Limit to <10% of daily calories"] | Sources: [e.g., Red meat, butter, coconut oil, cheese]
        - **Dietary Cholesterol:** [Provide mg/day, e.g., <300 mg] | Sources: [e.g., Egg yolks, shellfish, organ meats, full-fat dairy]

**Micronutrients:**
    **1. Vitamins:**
        - **Vitamin A (retinol/carotenoids):** [Provide mcg RAE/day, e.g., 700-900 mcg] | Sources: [e.g., Carrots, sweet potatoes, spinach, liver]
        - **Vitamin B Complex:**
            - B1 (Thiamine): [Provide mg/day, e.g., 1.1-1.2 mg] | Sources: [e.g., Pork, whole grains, nutritional yeast, black beans]
            - B2 (Riboflavin): [Provide mg/day, e.g., 1.1-1.3 mg] | Sources: [e.g., Dairy products, almonds, lean meat, mushrooms]
            - B3 (Niacin): [Provide mg NE/day, e.g., 14-16 mg] | Sources: [e.g., Chicken, tuna, peanuts, brown rice]
            - B5 (Pantothenic Acid): [Provide mg/day, e.g., 5 mg] | Sources: [e.g., Avocado, shiitake mushrooms, sunflower seeds, chicken]
            - B6 (Pyridoxine): [Provide mg/day, e.g., 1.3-1.7 mg] | Sources: [e.g., Chickpeas, salmon, potatoes, bananas]
            - B7 (Biotin): [Provide mcg/day, e.g., 30 mcg] | Sources: [e.g., Egg yolk, liver, salmon, avocado, nuts, seeds]
            - B9 (Folate): [Provide mcg DFE/day, e.g., 400 mcg] | Sources: [e.g., Leafy greens, lentils, beans, fortified grains]
            - B12 (Cobalamin): [Provide mcg/day, e.g., 2.4 mcg] | Sources: [e.g., Clams, tuna, beef, fortified nutritional yeast (for vegans)]
        - **Vitamin C (ascorbic acid):** [Provide mg/day, e.g., 75-90 mg] | Sources: [e.g., Bell peppers, oranges, broccoli, strawberries]
        - **Vitamin D (calciferol):** [Provide IU/day, e.g., 600-800 IU] | Sources: [e.g., Fatty fish (salmon), fortified milk, sunlight]
        - **Vitamin E (tocopherol):** [Provide mg/day, e.g., 15 mg] | Sources: [e.g., Almonds, sunflower seeds, spinach, avocado]
        - **Vitamin K (phylloquinone/menaquinone):** [Provide mcg/day, e.g., 90-120 mcg] | Sources: [e.g., Kale, spinach, broccoli, natto]

    **2. Minerals:**
        - **Calcium:** [Provide mg/day, e.g., 1000-1300 mg] | Sources: [e.g., Dairy, tofu, sardines, fortified plant milks]
        - **Phosphorus:** [Provide mg/day, e.g., 700 mg] | Sources: [e.g., Meat, fish, dairy, nuts, seeds, whole grains]
        - **Potassium:** [Provide mg/day, e.g., 2600-3400 mg] | Sources: [e.g., Bananas, potatoes, spinach, beans, lentils]
        - **Magnesium:** [Provide mg/day, e.g., 310-420 mg] | Sources: [e.g., Nuts, seeds, leafy greens, dark chocolate]
        - **Iron:** [Provide mg/day, e.g., 8-18 mg] | Sources: [e.g., Red meat, beans, lentils, spinach, fortified cereals]
        - **Zinc:** [Provide mg/day, e.g., 8-11 mg] | Sources: [e.g., Oysters, red meat, chickpeas, pumpkin seeds]
        - **Sodium:** [Provide mg/day, e.g., < 2300 mg] | Sources: [e.g., Limit processed foods; small amounts in vegetables]
        - **Chloride:** [Provide g/day, e.g., 1.8-2.3 g] | Sources: [e.g., Table salt, seaweed, tomatoes, celery]
        - **Trace Elements:**
            - Copper: [Provide mcg/day, e.g., 900 mcg] | Sources: [e.g., Oysters, shiitake mushrooms, cashews, seeds]
            - Manganese: [Provide mg/day, e.g., 1.8-2.3 mg] | Sources: [e.g., Mussels, whole grains, nuts, leafy vegetables]
            - Selenium: [Provide mcg/day, e.g., 55 mcg] | Sources: [e.g., Brazil nuts, seafood, organ meats, eggs]
            - Iodine: [Provide mcg/day, e.g., 150 mcg] | Sources: [e.g., Seaweed, cod, dairy, iodized salt]
            - Chromium: [Provide mcg/day, e.g., 25-35 mcg] | Sources: [e.g., Broccoli, grape juice, whole wheat products]
            - Molybdenum: [Provide mcg/day, e.g., 45 mcg] | Sources: [e.g., Legumes, grains, nuts, leafy vegetables]
            - Fluoride: [Provide mg/day, e.g., 3-4 mg] | Sources: [e.g., Fluoridated water, tea, grapes, raisins]

**Other Key Compounds:**
    - **Water:** [Provide L/day and glasses, e.g., 2.7-3.7 L (about 8-12 glasses)] | Tip: [Add a tip, e.g., "Drink more if active or in a hot climate."]
    - **Fiber:** [Provide g/day, e.g., 25-38 g] | Sources: [e.g., Oats, beans, apples, broccoli, psyllium husk]
    - **Electrolytes:** [Provide context, e.g., "Crucial for hydration; obtain from a balanced diet."] | Key Sources: [e.g., Sodium, Potassium, Magnesium from diet]
    - **Phytochemicals:**
        - **Polyphenols:** [Provide context, e.g., "Potent antioxidants."] | Sources: [e.g., Berries, dark chocolate, tea, coffee, red wine]
        - **Carotenoids:** [Provide context, e.g., "Protect cells from damage."] | Sources: [e.g., Carrots, tomatoes, spinach, kale]

**Actionable Advice & Recommendations:**
    - **General Tips:**
        - [Provide a tip about diet diversity.]
        - [Provide a tip about mindful eating and portion control.]
        - [Provide a tip about limiting processed foods, added sugars, and unhealthy fats.]
    - **Health-Specific Guidance:**
        # If no health condition is provided, output: "No specific health conditions were listed. Focus on general wellness by maintaining a balanced diet, staying hydrated, and engaging in regular physical activity."
        # If a health condition is provided, follow the example format below.
        # EXAMPLE START
        # - **Diabetes:**
        #   **Advice:** Prioritize complex carbohydrates with a low glycemic index (e.g., whole grains, legumes) to ensure stable blood sugar. Pair carbs with protein and healthy fats.
        #   **Key Nutrients/Deficiencies:** Monitor **Magnesium** and **Chromium**, which play roles in glucose metabolism. **B-Vitamins**, especially B12, can be impacted by some medications.
        # - **Hypertension:**
        #   **Advice:** Dramatically reduce sodium intake by avoiding processed foods. Increase intake of potassium-rich foods (fruits, vegetables) to counterbalance sodium's effect on blood pressure.
        #   **Key Nutrients/Deficiencies:** Focus on **Potassium, Calcium, and Magnesium**, as they are vital for blood pressure regulation.
        # EXAMPLE END
    - **Dietary Preference Tips:**
        - [Provide a specific, actionable tip based on the user's diet, e.g., For a Vegan: "Ensure reliable Vitamin B12 intake through fortified foods (nutritional yeast, plant milks) or a supplement, as it's not naturally present in plant foods."]
"""

SYSTEM_INSTRUCTION_FOOD_COMPARISON = """
Act as a concise nutritionist. Compare the two foods provided by the user.

**CRITICAL: Your entire response must be ONLY the HTML `<table>` element and nothing else. Do not include `<html>`, `<body>`, `<!DOCTYPE>`, markdown fences (```html), or any explanatory text before or after the table.**

**Output Instructions:**
1.  Generate a **single, complete HTML `<table>`**.
2.  The table must have a `<thead>` with the first column header as "Nutritional Metric (per 100g)" and subsequent headers as the food names.
3.  The `<tbody>` must contain one `<tr>` for each of the following metrics:
    - Calories (kcal)
    - Protein (g)
    - Carbohydrates (g)
    - Fiber (g)
    - Sugars (g)
    - Fat (g)
    - Key Vitamin
    - Key Mineral
    - Best For
4.  Provide concise, data-driven comparisons for each metric in its respective `<td>`.
5.  Highlight the "winner" for each metric (e.g., higher protein, lower sugar) with a simple emoji like ✅ or a brief comment.
"""

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/get_nutrient_recommendations', methods=['POST'])
def nutrient_recommendations():
    data = request.get_json()
    if not data:
        return jsonify({"error": "Invalid request body"}), 400

    prompt_data = {
        "age": data.get("age", "N/A"),
        "gender": data.get("gender", "N/A"),
        "height": data.get("height", "N/A"),
        "weight": data.get("weight", "N/A"),
        "activity_level": data.get("activity_level", "N/A"),
        "pregnancy_or_lactation": data.get("pregnancy_or_lactation", "None"),
        "health_condition": data.get("health_condition", "None"),
        "dietary_preferences": data.get("dietary_preferences", "N/A"),
    }

    prompt_text = "\n".join([f"{key.replace('_', ' ').title()}: {value}" for key, value in prompt_data.items()])

    try:
        contents = [types.Content(role="user", parts=[types.Part.from_text(text=prompt_text)])]

        config = types.GenerateContentConfig(
            system_instruction=SYSTEM_INSTRUCTION_NUTRITIONIST,
            temperature=0.4,
            response_mime_type="text/plain",
        )

        response_text = ""
        for chunk in client.models.generate_content_stream(
            model="gemini-2.0-flash-lite",
            contents=contents,
            config=config,
        ):
            if chunk.text:
                response_text += chunk.text

        if not response_text:
            raise ValueError("AI returned an empty response.")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"NutriAI_Report_{timestamp}"

        text_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{base_filename}.txt")
        with open(text_file_path, "w", encoding='utf-8') as f:
            f.write(response_text)

        formatted_html = format_recommendations_to_html(response_text)

        return jsonify({
            'recommendations': formatted_html,
            'download_token': base_filename
        })

    except Exception as e:
        logging.error(f"Error in recommendation endpoint: {e}", exc_info=True)
        return jsonify({"error": f"An internal server error occurred: {e}"}), 500

@app.route('/compare_foods', methods=['POST'])
def compare_foods():
    data = request.get_json()
    if not data or 'foods' not in data or len(data['foods']) != 2:
        return jsonify({"error": "Please provide exactly two foods to compare."}), 400

    foods = data['foods']
    prompt_text = f"Compare {foods[0]} and {foods[1]}"

    try:
        contents = [types.Content(role="user", parts=[types.Part.from_text(text=prompt_text)])]

        config = types.GenerateContentConfig(
            system_instruction=SYSTEM_INSTRUCTION_FOOD_COMPARISON,
            temperature=0.3,
            response_mime_type="text/plain",
        )

        raw_html_table = ""
        for chunk in client.models.generate_content_stream(
            model="gemini-2.0-flash-lite",
            contents=contents,
            config=config,
        ):
            if chunk.text:
                raw_html_table += chunk.text

        clean_html_table = format_comparison(raw_html_table)

        return jsonify({"comparison": clean_html_table})
    except Exception as e:
        logging.error(f"Error in food comparison endpoint: {e}", exc_info=True)
        return jsonify({"error": f"An internal server error occurred: {e}"}), 500

def _convert_txt_to_docx(txt_file_path):
    """Converts a text file to a DOCX document with specific A4 landscape formatting."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = 1
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    with open(txt_file_path, 'r', encoding='utf-8') as file:
        text_content = file.read().replace('**', '')
        for line in text_content.split('\n'):
            p = doc.add_paragraph(line.rstrip('\n'))
            p.paragraph_format.space_before = 0
            p.paragraph_format.space_after = 0
    docx_file_path = txt_file_path.replace('.txt', '.docx')
    doc.save(docx_file_path)
    logging.info(f"Successfully converted TXT to DOCX: {docx_file_path}")
    return docx_file_path

class PDF(FPDF):
    """Custom PDF class to remove default header and footer."""
    def header(self): pass
    def footer(self): pass

def _convert_docx_to_pdf(docx_file_path):
    import pyphen
    dic = pyphen.Pyphen(lang='en_US')

    pdf_file_path = docx_file_path.replace(".docx", ".pdf")
    doc = Document(docx_file_path)
    pdf = PDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()

    font_path = os.path.join(os.path.dirname(__file__), 'static', 'fonts', 'DejaVuSerif.ttf')
    if os.path.isfile(font_path):
        pdf.add_font('DejaVuSerif', '', font_path, uni=True)
        pdf.set_font('DejaVuSerif', size=12)
    else:
        pdf.set_font('Arial', size=12)

    base_left = pdf.l_margin
    max_width = pdf.w - pdf.r_margin

    for paragraph in doc.paragraphs:
        text = paragraph.text or ""
        if not text.strip():
            pdf.ln(5)
            continue

        # Calculate indentation (in mm)
        indent_pts = paragraph.paragraph_format.left_indent
        if indent_pts:
            indent_mm = indent_pts.pt * 0.3528  # convert points to mm
        else:
            indent_mm = 0

        # Set new x position for this paragraph
        pdf.set_x(base_left + indent_mm)
        usable_width = max_width - (base_left + indent_mm)

        # Hyphenation-aware wrapping
        words = text.split()
        line = ''
        for word in words:
            test_line = f"{line} {word}".strip()
            if pdf.get_string_width(test_line) <= usable_width:
                line = test_line
            else:
                # Word too long, try hyphenation
                hyphenated = dic.inserted(word, hyphen='-').split('-')
                for i, part in enumerate(hyphenated):
                    part += '-' if i < len(hyphenated) - 1 else ''
                    test_line = f"{line} {part}".strip()
                    if pdf.get_string_width(test_line) <= usable_width:
                        line = test_line
                    else:
                        pdf.set_x(base_left + indent_mm)
                        pdf.cell(0, 5, line)
                        pdf.ln(5)
                        line = part
        if line:
            pdf.set_x(base_left + indent_mm)
            pdf.cell(0, 5, line)
            pdf.ln(5)

    pdf.output(pdf_file_path)
    logging.info(f"Successfully converted DOCX to PDF with indentations preserved: {pdf_file_path}")
    return pdf_file_path

@app.route('/download/<token>')
def download_file(token):
    """Finds the report by its token, converts it to PDF via DOCX, and serves it."""
    sanitized_token = secure_filename(token)
    text_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{sanitized_token}.txt")
    if not os.path.exists(text_file_path):
        logging.warning(f"Download request for non-existent token: {token}")
        abort(404, description="Report not found or has expired.")
    try:
        docx_path = _convert_txt_to_docx(text_file_path)
        pdf_path = _convert_docx_to_pdf(docx_path)
        return send_file(pdf_path, as_attachment=True, download_name=f"{sanitized_token}.pdf")
    except Exception as e:
        logging.error(f"Failed to generate or send PDF for token {token}: {e}", exc_info=True)
        abort(500, description="An error occurred while generating the PDF report.")

def format_recommendations_to_html(text: str) -> str:
    """
    FIXED: A robust, single-pass parser to convert the AI's structured text into clean HTML.
    It correctly identifies top-level cards and all nested content, including special handling
    for macronutrient-style subheadings.
    """
    html_output = ""
    lines = text.strip().split('\n')

    icon_map = {
        "BMI": "fa-weight", "ESTIMATED DAILY CALORIES": "fa-fire",
        "MACRONUTRIENTS": "fa-pizza-slice", "MICRONUTRIENTS": "fa-pills",
        "OTHER KEY COMPOUNDS": "fa-tint", "ACTIONABLE ADVICE & RECOMMENDATIONS": "fa-clipboard-check"
    }

    in_card = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        indent_level = len(line) - len(line.lstrip(' '))
        is_card_header = stripped.startswith('**') and stripped.endswith('**') and ':' in stripped and indent_level == 0

        if is_card_header:
            if in_card:
                html_output += "</div></div>\n"
            parts = stripped.replace('**', '').split(':', 1)
            title, content_on_same_line = parts[0].strip(), parts[1].strip()
            card_id, icon = title.lower().replace(' ', '_').replace('&', 'and'), icon_map.get(title.upper(), "fa-info-circle")
            html_output += f"""<div class='result-card' id='{card_id}'>
    <button class='result-card-header' onclick='app.toggleCardBody("{card_id}")'>
        <span><i class='fas {icon}'></i> {title}</span>
        <i class='fas fa-chevron-down card-chevron'></i>
    </button>
    <div class='result-card-body'>"""
            in_card = True
            if content_on_same_line:
                html_output += f"<p class='main-value'>{content_on_same_line}</p>\n"
            continue

        if not in_card:
            continue

        # This new condition specifically finds indented, bolded, numbered list items with a colon,
        # which is the pattern for Macronutrient subheadings (e.g., "**1. Carbohydrates:** value").
        is_macro_style_header = stripped.startswith('**') and re.match(r'\*\*\d\.', stripped) and ':' in stripped

        if is_macro_style_header:
            clean_line = stripped.replace('**', '')
            parts = clean_line.split(':', 1)
            header_text = parts[0].strip() + ':'
            value_text = parts[1].strip() if len(parts) > 1 else ""

            # Create the h3 tag for the header part
            html_output += f"<h3 class='section-heading' style='margin-left: {indent_level * 2}px;'>{header_text}</h3>\n"

            # If there was a value on the same line, process and display it
            if value_text:
                if '|' in value_text:
                    value_parts = value_text.split('|', 1)
                    html_output += f"""<div class='nutrient-item' style='margin-left: {indent_level * 2 + 10}px;'>
    <span class='nutrient-name'>{value_parts[0].strip()}</span>
    <span class='nutrient-source'>{value_parts[1].strip()}</span>
</div>\n"""
                else:
                    html_output += f"<p style='margin-left: {indent_level * 2 + 10}px;'>{value_text}</p>\n"
            continue # Important: We've handled this line, so skip to the next one

        if stripped.startswith('**') and stripped.endswith('**') and ':' in stripped:
            tag = "h3" if indent_level < 8 else "h4"
            html_output += f"<{tag} class='section-heading' style='margin-left: {indent_level * 2}px;'>{stripped.replace('**', '')}</{tag}>\n"
        elif stripped.startswith('-'):
            item_text = stripped.replace('**', '')[1:].strip()
            if '|' in item_text:
                parts = item_text.split('|', 1)
                source_part = parts[1].strip().replace('Sources:', '').replace('Tip:', '').strip()
                html_output += f"""<div class='nutrient-item' style='margin-left: {indent_level * 2 + 10}px;'>
    <span class='nutrient-name'>{parts[0].strip()}</span>
    <span class='nutrient-source'>{source_part}</span>
</div>\n"""
            else:
                html_output += f"<p class='list-item' style='margin-left: {indent_level * 2}px;'>• {item_text}</p>\n"
        elif '|' in stripped:
            parts = stripped.replace('**', '').split('|', 1)
            source_part = parts[1].strip().replace('Sources:', '').replace('Tip:', '').strip()
            html_output += f"""<div class='nutrient-item' style='margin-left: {indent_level * 2}px;'>
    <span class='nutrient-name'>{parts[0].strip()}</span>
    <span class='nutrient-source'>{source_part}</span>
</div>\n"""
        else:
            html_output += f"<p style='margin-left: {indent_level * 2}px;'>{stripped.replace('**', '')}</p>\n"

    if in_card:
        html_output += "</div></div>\n"

    return html_output


def format_comparison(text: str) -> str:
    """
    Cleans the model's output to extract just the HTML table for crop comparison.
    Handles cases where the model might still include markdown fences or explanatory text.
    """
    # Remove markdown fences and surrounding whitespace
    cleaned_text = text.strip().removeprefix('```html').removesuffix('```').strip()

    # Use regex to find the table, which is more robust
    table_match = re.search(r'(<table.*?>.*?</table\s*>)', cleaned_text, re.DOTALL | re.IGNORECASE)

    if table_match:
        # If a table is found, return it directly. This is the ideal case.
        return table_match.group(1)
    else:
        # Fallback for unexpected format: return a formatted error.
        logging.warning(f"Comparison format error. AI output was: {text}")
        return f"""
        <div class="no-data error">
            <i class="fas fa-exclamation-triangle"></i>
            <h3>Comparison Error</h3>
            <p>The AI model returned data in an unexpected format. Please try your query again.</p>
        </div>
        """


if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=5000)